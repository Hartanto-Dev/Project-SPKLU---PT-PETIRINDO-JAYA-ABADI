import os
import shutil
import docx
from docx.shared import Inches
from PIL import Image

def convert_to_standard_jpg(input_path, output_path):
    try:
        with Image.open(input_path) as img:
            img.convert("RGB").save(output_path, "JPEG")
        print(f"Successfully converted {input_path} to standard JPEG at {output_path}")
        return True
    except Exception as e:
        print(f"Error converting {input_path}: {e}")
        return False

def update_skripsi():
    # We will copy the original docx again to clean up any partial changes from the failed run
    backup_file = "Hartanto - Skripsi - Demo - Backup.docx"
    src_file = "Hartanto - Skripsi - Demo.docx"
    
    if os.path.exists(backup_file):
        shutil.copyfile(backup_file, src_file)
        print("Restored original document from backup to start fresh.")
    else:
        shutil.copyfile(src_file, backup_file)
        print(f"Backup created at {backup_file}")
    
    # Pre-convert potential problematic JPEGs to standard formats
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/Gambar 1.1. Siklus Hidup Pengembangan Sistem Informasi.jpeg", 
        "Daftar Gambar Skripsi/SiklusHidup_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/proses scrum.jpeg", 
        "Daftar Gambar Skripsi/Scrum_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/SPKLU.jpg", 
        "Daftar Gambar Skripsi/SPKLU_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/Tahapan Penelitian.jpg", 
        "Daftar Gambar Skripsi/Tahapan_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/Ranvangan Penelitian.jpg", 
        "Daftar Gambar Skripsi/Rancangan_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/Activity Diagram.jpeg", 
        "Daftar Gambar Skripsi/Activity_Standard.jpg"
    )
    convert_to_standard_jpg(
        "Daftar Gambar Skripsi/Diagram ERD.jpg", 
        "Daftar Gambar Skripsi/ERD_Standard.jpg"
    )
    
    # Pre-convert screenshots just to be safe
    convert_to_standard_jpg("Screenshot Page Home.jpg", "Screenshot_Home_Standard.jpg")
    convert_to_standard_jpg("Screenshot Page Dashboard.jpg", "Screenshot_Dashboard_Standard.jpg")
    convert_to_standard_jpg("Screenshot About Us.jpg", "Screenshot_AboutUs_Standard.jpg")
    convert_to_standard_jpg("Screenshot Contacts.jpg", "Screenshot_Contacts_Standard.jpg")
    convert_to_standard_jpg("Mobile Hamburger.jpg", "Mobile_Hamburger_Standard.jpg")

    doc = docx.Document(src_file)
    
    # Mapping for placeholders using original captions
    # Key: original caption prefix, Value: path to standard image
    image_mapping = {
        "Gambar 2.1": "Daftar Gambar Skripsi/SiklusHidup_Standard.jpg",
        "Gambar 2.2": "Daftar Gambar Skripsi/visualisasi data.png",
        "Gambar 2.3": "Daftar Gambar Skripsi/Scrum_Standard.jpg",
        "Gambar 2.4": "Daftar Gambar Skripsi/SPKLU_Standard.jpg",
        "Gambar 3.1": "Daftar Gambar Skripsi/Tahapan_Standard.jpg",
        "Gambar 3.2": "Daftar Gambar Skripsi/Rancangan_Standard.jpg",
        "Gambar 4.1": "Screenshot_Home_Standard.jpg",
        "Gambar 4.2": "Screenshot_Dashboard_Standard.jpg",
        "Gambar 4.3": "Screenshot_AboutUs_Standard.jpg",
        "Gambar 4.4": "Screenshot_Contacts_Standard.jpg",
        "Gambar 4.5": "Mobile_Hamburger_Standard.jpg"
    }

    # 1. Replace the [Sisipkan gambar di sini] placeholders first using original captions
    print("Replacing image placeholders...")
    placeholders_indices = []
    for i, p in enumerate(doc.paragraphs):
        if "[Sisipkan gambar di sini]" in p.text:
            placeholders_indices.append(i)
            
    for idx in sorted(placeholders_indices, reverse=True):
        caption_para = doc.paragraphs[idx + 1]
        caption_text = caption_para.text
        
        matched_img = None
        for key, img_path in image_mapping.items():
            if key in caption_text:
                matched_img = img_path
                break

        if matched_img and os.path.exists(matched_img):
            p_placeholder = doc.paragraphs[idx]
            p_placeholder.text = ""
            run = p_placeholder.add_run()
            
            width_inch = 5.5 if "Hamburger" not in matched_img else 2.5
            run.add_picture(matched_img, width=Inches(width_inch))
            p_placeholder.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            caption_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            print(f"Inserted {matched_img} for caption: '{caption_text[:50]}...'")
        else:
            print(f"Warning: No matching image file found for caption '{caption_text[:50]}...' (path checked: {matched_img})")

    # 2. Shift the existing figure numbers in Chapter 4 captions (from 4.1-4.5 to 4.5-4.9)
    print("Shifting caption numbers...")
    for p in doc.paragraphs:
        if "Gambar 4.1" in p.text and "Tampilan Halaman Home" in p.text:
            p.text = p.text.replace("Gambar 4.1", "Gambar 4.5")
        elif "Gambar 4.2" in p.text and "Tampilan Halaman Dashboard" in p.text:
            p.text = p.text.replace("Gambar 4.2", "Gambar 4.6")
        elif "Gambar 4.3" in p.text and "Tampilan Halaman About Us" in p.text:
            p.text = p.text.replace("Gambar 4.3", "Gambar 4.7")
        elif "Gambar 4.4" in p.text and "Tampilan Halaman Contacts" in p.text:
            p.text = p.text.replace("Gambar 4.4", "Gambar 4.8")
        elif "Gambar 4.5" in p.text and "Tampilan Navbar" in p.text:
            p.text = p.text.replace("Gambar 4.5", "Gambar 4.9")

    # 3. Insert the new UML diagrams and descriptions
    print("Inserting UML diagrams and descriptions...")
    
    # Locate 4.1.1, 4.1.2, 4.1.3 paragraphs again
    p_411_idx = -1
    p_412_idx = -1
    p_413_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("4.1.1"):
            p_411_idx = i
        elif p.text.startswith("4.1.2"):
            p_412_idx = i
        elif p.text.startswith("4.1.3"):
            p_413_idx = i

    # --- INSERT 3: Activity Diagram under 4.1.3 Alur Navigasi Pengguna ---
    if p_413_idx != -1:
        insert_idx = p_413_idx
        for j in range(p_413_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.startswith("4.1.4") or doc.paragraphs[j].text.startswith("4.2"):
                insert_idx = j - 1
                break
        
        target_p = doc.paragraphs[insert_idx + 1]
        
        desc_text = (
            "Alur kerja interaktif (workflow) yang melibatkan kolaborasi antara Admin, Sistem SPKLU, dan User "
            "dalam lingkungan platform digambarkan secara terstruktur pada Gambar 4.3. Aktivitas dimulai (start state) "
            "dari sisi Admin yang melakukan Login Sistem. Sistem SPKLU kemudian memvalidasi kredensial tersebut dan "
            "menampilkan halaman utama administrasi. Setelah berhasil masuk ke dalam sistem, Admin dapat melaksanakan "
            "tiga fungsi pemantauan dan pengelolaan utama: (1) Pengelolaan Data SPKLU di mana Admin memperbarui data, "
            "lalu sistem memproses penyimpanan data tersebut sehingga dapat diakses oleh User; (2) Visualisasi Data di "
            "mana Admin mengakses dasbor perkembangan yang memicu sistem untuk mengolah data statistik dan menampilkannya; "
            "serta (3) Monitoring SPKLU di mana Admin memantau data operasional secara real-time. Di sisi lain, User dapat "
            "berinteraksi dengan peta untuk melihat detail informasi SPKLU serta melakukan pencarian atau filter tipe pengisian daya. "
            "Sistem memproses parameter tersebut untuk menampilkan hasil spasial yang relevan. Alur aktivitas ini diakhiri dengan "
            "proses pengambilan keputusan (decision node) oleh Admin untuk melakukan fungsi Logout Sistem."
        )
        
        p_desc = target_p.insert_paragraph_before(desc_text)
        p_desc.style = 'Normal'
        
        p_img = target_p.insert_paragraph_before()
        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture("Daftar Gambar Skripsi/Activity_Standard.jpg", width=Inches(5.0))
        
        p_cap = target_p.insert_paragraph_before("Gambar 4.3 Activity Diagram Platform Web Petirindo SPKLU")
        p_cap.style = 'Normal'
        p_cap.bold = True
        p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        print("Inserted Activity Diagram at 4.1.3")

    # --- INSERT 2: Architecture & ERD Diagram under 4.1.2 Arsitektur Sistem ---
    if p_412_idx != -1:
        insert_idx = p_412_idx
        for j in range(p_412_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.startswith("4.1.3"):
                insert_idx = j - 1
                break
                
        target_p = doc.paragraphs[insert_idx + 1]
        
        # --- ERD ---
        erd_desc = (
            "Implementasi penyimpanan data pada Platform Web Petirindo SPKLU dirancang menggunakan skema relasi logis "
            "atau Entity Relationship Diagram (ERD) seperti yang ditunjukkan pada Gambar 4.2. Mengingat sistem ini dikembangkan "
            "dengan pendekatan Frontend-Only Static Web App, struktur data logis ini diwujudkan melalui objek data terstruktur "
            "pada berkas spkluData.js dan objek konteks pengguna (UserContext) yang saling terhubung. Rancangan ini terdiri dari "
            "dua entitas utama, yaitu USERS (menyimpan data pengguna administratif) dan SPKLU (menyimpan detail informasi stasiun "
            "pengisian daya). Hubungan relasi antara entitas USERS dan SPKLU adalah Satu-ke-Banyak (One-to-Many atau 1:N) melalui "
            "kunci tamu created_by pada entitas SPKLU."
        )
        p_erd_desc = target_p.insert_paragraph_before(erd_desc)
        p_erd_desc.style = 'Normal'
        
        p_erd_img = target_p.insert_paragraph_before()
        p_erd_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = p_erd_img.add_run()
        run.add_picture("Daftar Gambar Skripsi/ERD_Standard.jpg", width=Inches(5.5))
        
        p_erd_cap = target_p.insert_paragraph_before("Gambar 4.2 Entity Relationship Diagram Platform Web Petirindo SPKLU")
        p_erd_cap.style = 'Normal'
        p_erd_cap.bold = True
        p_erd_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        # Spacer
        target_p.insert_paragraph_before()
        
        # --- Architecture Diagram ---
        arch_desc = (
            "Secara visual, hubungan interaksi antar-komponen dalam arsitektur Platform Web Petirindo SPKLU digambarkan "
            "pada Gambar 4.1. Diagram ini memetakan bagaimana Browser Pengguna berinteraksi dengan Web Server untuk mengunduh "
            "berkas bundle statis secara asinkron, membaca data SPKLU statis dari berkas spkluData.js, serta melakukan "
            "request ubin peta dasar (basemap) ke Server Ubin OpenStreetMap secara asinkron."
        )
        p_arch_desc = target_p.insert_paragraph_before(arch_desc)
        p_arch_desc.style = 'Normal'
        
        p_arch_img = target_p.insert_paragraph_before()
        p_arch_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = p_arch_img.add_run()
        run.add_picture("Daftar Gambar Skripsi/Arsitektur Diagram.png", width=Inches(5.5))
        
        p_arch_cap = target_p.insert_paragraph_before("Gambar 4.1 Diagram Arsitektur Platform Web Petirindo SPKLU")
        p_arch_cap.style = 'Normal'
        p_arch_cap.bold = True
        p_arch_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        print("Inserted Architecture & ERD Diagrams at 4.1.2")

    # --- INSERT 1: Use Case Diagram under 4.1.1 Gambaran Umum Sistem ---
    if p_411_idx != -1:
        insert_idx = p_411_idx
        for j in range(p_411_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[j].text.startswith("4.1.2"):
                insert_idx = j - 1
                break
                
        target_p = doc.paragraphs[insert_idx + 1]
        
        uc_desc = (
            "Fungsionalitas serta batasan interaksi pengguna di dalam Platform Web Petirindo SPKLU didefinisikan secara spesifik "
            "melalui use case diagram pada Gambar 4.4. Di dalam sistem ini, terdapat 3 (tiga) aktor utama yang memiliki hak akses "
            "berbeda, yaitu User (Pengguna Umum) yang dapat melihat peta, pencarian, dan detail informasi SPKLU; Admin yang dapat "
            "mengelola data SPKLU dan melihat dasbor monitoring; serta Manajemen yang dapat melihat dasbor visualisasi data tanpa "
            "memiliki hak manipulasi data."
        )
        p_uc_desc = target_p.insert_paragraph_before(uc_desc)
        p_uc_desc.style = 'Normal'
        
        p_uc_img = target_p.insert_paragraph_before()
        p_uc_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = p_uc_img.add_run()
        run.add_picture("Daftar Gambar Skripsi/Use Case Diagram.png", width=Inches(5.2))
        
        p_uc_cap = target_p.insert_paragraph_before("Gambar 4.4 Use Case Diagram Platform Web Petirindo SPKLU")
        p_uc_cap.style = 'Normal'
        p_uc_cap.bold = True
        p_uc_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        
        print("Inserted Use Case Diagram at 4.1.1")
        
    doc.save(src_file)
    print("Success: Skripsi document updated with all diagrams, captions, and descriptions!")

    # Cleanup temporary pre-converted files
    temp_files = [
        "Daftar Gambar Skripsi/SiklusHidup_Standard.jpg",
        "Daftar Gambar Skripsi/Scrum_Standard.jpg",
        "Daftar Gambar Skripsi/SPKLU_Standard.jpg",
        "Daftar Gambar Skripsi/Tahapan_Standard.jpg",
        "Daftar Gambar Skripsi/Rancangan_Standard.jpg",
        "Daftar Gambar Skripsi/Activity_Standard.jpg",
        "Daftar Gambar Skripsi/ERD_Standard.jpg",
        "Screenshot_Home_Standard.jpg",
        "Screenshot_Dashboard_Standard.jpg",
        "Screenshot_AboutUs_Standard.jpg",
        "Screenshot_Contacts_Standard.jpg",
        "Mobile_Hamburger_Standard.jpg"
    ]
    for f in temp_files:
        if os.path.exists(f):
            os.remove(f)

if __name__ == '__main__':
    update_skripsi()
